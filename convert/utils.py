from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import io
import PyPDF2
from docx import Document

def convert_word_to_pdf(word_file):
    # Load the Word file
    doc = Document(word_file)
    pdf_buffer = io.BytesIO()
    
    c = canvas.Canvas(pdf_buffer, pagesize=letter)
    text_object = c.beginText(40, 750)
    text_object.setFont("Helvetica", 12)
    
    for para in doc.paragraphs:
        text_object.textLine(para.text)
    
    c.drawText(text_object)
    c.showPage()
    c.save()
    
    pdf_buffer.seek(0)
    return pdf_buffer

def convert_pdf_to_word(pdf_file):
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    doc = Document()

    for page_num in range(len(pdf_reader.pages)):
        page = pdf_reader.pages[page_num]
        text = page.extract_text()
        doc.add_paragraph(text)

    word_buffer = io.BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)
    return word_buffer
